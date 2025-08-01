from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.db.models import Count, Q
from django.utils import timezone
from django.conf import settings
from django.core.files.storage import default_storage
from django.contrib.auth.decorators import login_required
from django.utils.dateparse import parse_datetime
from datetime import datetime, timedelta
import json
from docx import Document
from io import BytesIO

from .models import Applicant, Interviewer, Department
from .forms import ApplicationForm, InterviewerLoginForm, DepartmentForm, InterviewerForm


# ------------------------------
# 公共视图
# ------------------------------

def homepage(request):
    """首页视图"""
    departments = Department.objects.all()
    return render(request, 'homepage.html', {
        'departments': departments
    })


def application_form(request):
    """新生报名表单页面"""
    if request.method == 'POST':
        form = ApplicationForm(request.POST, request.FILES)
        if form.is_valid():
            applicant = form.save()

            # 自动将简历添加到第一志愿部门面试官的待筛选列表
            first_department = applicant.first_choice
            if first_department:
                # 获取该部门的面试官
                interviewers = Interviewer.objects.filter(department_team=first_department)
                if interviewers.exists():
                    # 简单分配给第一个面试官（实际应用中可按负载均衡）
                    interviewer = interviewers.first()
                    if str(applicant.id) not in interviewer.resumes_to_screen:
                        interviewer.resumes_to_screen.append(str(applicant.id))
                        interviewer.save()

            messages.success(request, "报名成功！我们将尽快处理您的申请。")
            return redirect('application_form')
        else:
            messages.error(request, "表单填写有误，请检查后重新提交。")
    else:
        form = ApplicationForm()

    departments = Department.objects.all()
    return render(request, 'application_form.html', {
        'form': form,
        'departments': departments
    })


# ------------------------------
# 面试官视图
# ------------------------------

def interviewer_login(request):
    """面试官登录"""
    if request.method == 'POST':
        form = InterviewerLoginForm(request.POST)
        if form.is_valid():
            student_id = form.cleaned_data['student_id']
            password = form.cleaned_data['password']

            try:
                interviewer = Interviewer.objects.get(student_id=student_id)
                if interviewer.check_password(password):
                    # 登录成功，存储会话
                    request.session['interviewer_id'] = str(interviewer.id)
                    return redirect('interviewer_dashboard')
                else:
                    messages.error(request, "密码错误")
            except Interviewer.DoesNotExist:
                messages.error(request, "该学号未注册为面试官")
    else:
        form = InterviewerLoginForm()

    return render(request, 'interviewer/login.html', {'form': form})


def interviewer_logout(request):
    """面试官登出"""
    if 'interviewer_id' in request.session:
        del request.session['interviewer_id']
    messages.success(request, "已成功登出")
    return redirect('interviewer_login')


def interviewer_dashboard(request):
    """面试官仪表盘"""
    if 'interviewer_id' not in request.session:
        return redirect('interviewer_login')

    interviewer = get_object_or_404(
        Interviewer,
        id=request.session['interviewer_id']
    )

    # 统计数据
    total_resumes = len(interviewer.resumes_to_screen)
    pending_interviews = interviewer.pending_interviews

    # 获取待筛选简历
    resume_ids = [id for id in interviewer.resumes_to_screen if id]
    resumes_to_screen = Applicant.objects.filter(id__in=resume_ids)

    # 获取已安排的面试
    scheduled_interviews = Applicant.objects.filter(
        Q(first_interviewer=interviewer, first_interview_status='scheduled') |
        Q(second_interviewer=interviewer, second_interview_status='scheduled')
    )

    return render(request, 'interviewer/dashboard.html', {
        'interviewer': interviewer,
        'total_resumes': total_resumes,
        'pending_interviews': pending_interviews,
        'resumes_to_screen': resumes_to_screen,
        'scheduled_interviews': scheduled_interviews
    })


def application_list(request):
    """申请列表"""
    if 'interviewer_id' not in request.session:
        return redirect('interviewer_login')

    interviewer = get_object_or_404(
        Interviewer,
        id=request.session['interviewer_id']
    )

    # 获取该面试官需要处理的申请
    resume_ids = [id for id in interviewer.resumes_to_screen if id]
    applications = Applicant.objects.filter(
        Q(id__in=resume_ids) |
        Q(first_interviewer=interviewer) |
        Q(second_interviewer=interviewer)
    ).distinct()

    return render(request, 'interviewer/application_list.html', {
        'applications': applications,
        'interviewer': interviewer
    })


def application_detail(request, applicant_id):
    """申请详情"""
    if 'interviewer_id' not in request.session:
        return redirect('interviewer_login')

    interviewer = get_object_or_404(
        Interviewer,
        id=request.session['interviewer_id']
    )
    applicant = get_object_or_404(Applicant, id=applicant_id)

    # 检查权限
    has_access = (
            str(applicant.id) in interviewer.resumes_to_screen or
            applicant.first_interviewer == interviewer or
            applicant.second_interviewer == interviewer
    )

    if not has_access:
        messages.error(request, "您没有权限查看该申请")
        return redirect('application_list')

    if request.method == 'POST':
        # 处理简历筛选结果
        if 'screening_result' in request.POST:
            result = request.POST.get('screening_result')
            applicant.resume_screening_result = result
            applicant.save()

            # 从待筛选列表中移除
            if str(applicant.id) in interviewer.resumes_to_screen:
                interviewer.resumes_to_screen.remove(str(applicant.id))
                interviewer.save()

            messages.success(request, "简历筛选已完成")
            return redirect('application_detail', applicant_id=applicant_id)

        # 处理面试结果
        if 'interview_result' in request.POST:
            result = request.POST.get('interview_result')
            feedback = request.POST.get('feedback', '')

            # 判断是第一志愿还是第二志愿面试
            if applicant.first_interviewer == interviewer:
                applicant.first_interview_result = result
                applicant.first_interview_status = 'completed'
            elif applicant.second_interviewer == interviewer:
                applicant.second_interview_result = result
                applicant.second_interview_status = 'completed'

            # 保存反馈
            if feedback:
                applicant.feedback = feedback

            # 更新面试官的待面试数
            if interviewer.pending_interviews > 0:
                interviewer.pending_interviews -= 1
                interviewer.save()

            applicant.save()
            messages.success(request, "面试结果已提交")
            return redirect('application_detail', applicant_id=applicant_id)

    return render(request, 'interviewer/application_detail.html', {
        'applicant': applicant,
        'interviewer': interviewer
    })


def interview_schedule(request):
    """面试安排"""
    if 'interviewer_id' not in request.session:
        return redirect('interviewer_login')

    interviewer = get_object_or_404(
        Interviewer,
        id=request.session['interviewer_id']
    )

    if request.method == 'POST':
        # 保存可用面试时间
        time_slots = json.loads(request.POST.get('time_slots', '{}'))
        interviewer.available_interview_times = time_slots
        interviewer.save()
        messages.success(request, "面试时间已更新")
        return redirect('interview_schedule')

    # 获取该面试官的所有面试
    interviews = Applicant.objects.filter(
        Q(first_interviewer=interviewer) |
        Q(second_interviewer=interviewer)
    ).filter(
        Q(first_interview_status__in=['scheduled', 'completed']) |
        Q(second_interview_status__in=['scheduled', 'completed'])
    )

    return render(request, 'interviewer/schedule.html', {
        'interviewer': interviewer,
        'interviews': interviews,
        'available_times': json.dumps(interviewer.available_interview_times)
    })


def time_slots(request):
    """时间槽管理"""
    if 'interviewer_id' not in request.session:
        return redirect('interviewer_login')

    interviewer = get_object_or_404(
        Interviewer,
        id=request.session['interviewer_id']
    )

    if request.method == 'POST':
        # 处理时间槽设置
        action = request.POST.get('action')

        if action == 'add_time_slot':
            time_str = request.POST.get('time')
            max_count = int(request.POST.get('max_count', 5))

            # 验证时间格式
            try:
                datetime.fromisoformat(time_str)
                interviewer.available_interview_times[time_str] = max_count
                interviewer.save()
                return JsonResponse({'status': 'success'})
            except ValueError:
                return JsonResponse({'status': 'error', 'message': '时间格式不正确'}, status=400)

        elif action == 'delete_time_slot':
            time_str = request.POST.get('time')
            if time_str in interviewer.available_interview_times:
                del interviewer.available_interview_times[time_str]
                interviewer.save()
                return JsonResponse({'status': 'success'})

    return render(request, 'interviewer/time_slots.html', {
        'interviewer': interviewer,
        'available_times': interviewer.available_interview_times
    })


# ------------------------------
# 管理员视图
# ------------------------------

def admin_required(view_func):
    """管理员权限装饰器"""

    def wrapper(request, *args, **kwargs):
        # 这里假设管理员是特殊的面试官或有单独的管理员模型
        # 实际应用中应根据您的权限系统进行修改
        if 'interviewer_id' not in request.session:
            return redirect('admin_login')

        interviewer = get_object_or_404(
            Interviewer,
            id=request.session['interviewer_id']
        )

        # 假设部门名称为"管理员"的是管理员
        if interviewer.department_team.name != "管理员":
            messages.error(request, "您没有管理员权限")
            return redirect('interviewer_dashboard')

        return view_func(request, *args, **kwargs)

    return wrapper


def admin_login(request):
    """管理员登录（复用面试官登录）"""
    if request.method == 'POST':
        form = InterviewerLoginForm(request.POST)
        if form.is_valid():
            student_id = form.cleaned_data['student_id']
            password = form.cleaned_data['password']

            try:
                interviewer = Interviewer.objects.get(student_id=student_id)
                # 验证密码和管理员身份
                if interviewer.check_password(password) and interviewer.department_team.name == "管理员":
                    request.session['interviewer_id'] = str(interviewer.id)
                    return redirect('admin_dashboard')
                elif interviewer.department_team.name != "管理员":
                    messages.error(request, "您没有管理员权限")
                else:
                    messages.error(request, "密码错误")
            except Interviewer.DoesNotExist:
                messages.error(request, "该学号未注册")
    else:
        form = InterviewerLoginForm()

    return render(request, 'admin/login.html', {'form': form})


@admin_required
def admin_dashboard(request):
    """管理员仪表盘"""
    # 统计数据
    total_applicants = Applicant.objects.count()
    resumes_screened = Applicant.objects.exclude(resume_screening_result='pending').count()
    interviewed = Applicant.objects.filter(
        Q(first_interview_status='completed') |
        Q(second_interview_status='completed')
    ).count()
    admitted = Applicant.objects.filter(admission_result=True).count()

    # 部门统计
    department_stats = []
    for dept in Department.objects.all():
        first_choice_count = Applicant.objects.filter(first_choice=dept).count()
        second_choice_count = Applicant.objects.filter(second_choice=dept).count()
        admitted_count = Applicant.objects.filter(admitted_department=dept).count()

        department_stats.append({
            'name': dept.name,
            'first_choice': first_choice_count,
            'second_choice': second_choice_count,
            'total': first_choice_count + second_choice_count,
            'admitted': admitted_count
        })

    # 时间趋势
    date_range = [timezone.now() - timedelta(days=i) for i in range(7)]
    date_labels = [date.strftime('%m-%d') for date in date_range[::-1]]
    daily_counts = []

    for date in date_range[::-1]:
        next_day = date + timedelta(days=1)
        count = Applicant.objects.filter(created_at__gte=date, created_at__lt=next_day).count()
        daily_counts.append(count)

    # 待裁决的重复录取
    duplicate_admissions = []
    applicants = Applicant.objects.all()
    for applicant in applicants:
        if (applicant.first_interview_result == 'passed' and
                applicant.second_interview_result == 'passed'):
            duplicate_admissions.append(applicant)

    return render(request, 'admin/dashboard.html', {
        'total_applicants': total_applicants,
        'resumes_screened': resumes_screened,
        'interviewed': interviewed,
        'admitted': admitted,
        'department_stats': department_stats,
        'date_labels': date_labels,
        'daily_counts': daily_counts,
        'duplicate_admissions': len(duplicate_admissions)
    })


@admin_required
def admin_departments(request):
    """部门管理"""
    departments = Department.objects.all()

    if request.method == 'POST':
        if 'add_department' in request.POST:
            form = DepartmentForm(request.POST)
            if form.is_valid():
                form.save()
                messages.success(request, "部门已添加")
                return redirect('admin_departments')
        elif 'delete_department' in request.POST:
            dept_id = request.POST.get('dept_id')
            dept = get_object_or_404(Department, id=dept_id)
            dept_name = dept.name
            dept.delete()
            messages.success(request, f"部门 {dept_name} 已删除")
            return redirect('admin_departments')

    form = DepartmentForm()
    return render(request, 'admin/departments.html', {
        'departments': departments,
        'form': form
    })


@admin_required
def admin_applications(request):
    """报名信息汇总"""
    # 支持筛选
    status_filter = request.GET.get('status', 'all')
    department_filter = request.GET.get('department', 'all')

    applications = Applicant.objects.all().order_by('-created_at')

    # 应用筛选
    if status_filter != 'all':
        applications = applications.filter(resume_screening_result=status_filter)

    if department_filter != 'all' and department_filter:
        applications = applications.filter(
            Q(first_choice__name=department_filter) |
            Q(second_choice__name=department_filter)
        )

    departments = Department.objects.all()
    status_choices = dict(Applicant.STATUS_CHOICES)

    return render(request, 'admin/applications.html', {
        'applications': applications,
        'departments': departments,
        'status_choices': status_choices,
        'current_status': status_filter,
        'current_department': department_filter
    })


@admin_required
def admin_decisions(request):
    """录取裁决（处理重复录取）"""
    # 获取需要裁决的申请者（同时通过两个志愿面试）
    applicants = Applicant.objects.filter(
        first_interview_result='passed',
        second_interview_result='passed'
    )

    if request.method == 'POST':
        applicant_id = request.POST.get('applicant_id')
        decision = request.POST.get('decision')
        applicant = get_object_or_404(Applicant, id=applicant_id)

        if decision == 'first':
            applicant.admission_result = True
            applicant.admitted_department = applicant.first_choice
        elif decision == 'second':
            applicant.admission_result = True
            applicant.admitted_department = applicant.second_choice
        elif decision == 'neither':
            applicant.admission_result = False
            applicant.admitted_department = None

        applicant.save()
        messages.success(request, "裁决已完成")
        return redirect('admin_decisions')

    return render(request, 'admin/decisions.html', {
        'applicants': applicants
    })


@admin_required
def admin_admissions(request):
    """录取名单"""
    department_filter = request.GET.get('department', 'all')

    admissions = Applicant.objects.filter(admission_result=True)

    if department_filter != 'all' and department_filter:
        admissions = admissions.filter(admitted_department__name=department_filter)

    departments = Department.objects.all()

    return render(request, 'admin/admissions.html', {
        'admissions': admissions,
        'departments': departments,
        'current_department': department_filter
    })


@admin_required
def export_admissions(request):
    """导出录取名单为Word文档"""
    department_filter = request.GET.get('department', 'all')

    admissions = Applicant.objects.filter(admission_result=True)

    if department_filter != 'all' and department_filter:
        admissions = admissions.filter(admitted_department__name=department_filter)
        dept_name = department_filter
    else:
        dept_name = "所有部门"

    # 创建Word文档
    doc = Document()
    doc.add_heading(f'拾光文创招新录取名单 - {dept_name}', 0)

    # 添加表格
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '姓名'
    hdr_cells[2].text = '学号'
    hdr_cells[3].text = '年级'
    hdr_cells[4].text = '专业'
    hdr_cells[5].text = '录取部门'

    # 填充数据
    for i, applicant in enumerate(admissions.order_by('admitted_department__name', 'name'), 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = applicant.name
        row_cells[2].text = applicant.student_id
        row_cells[3].text = dict(Applicant.GRADE_CHOICES)[applicant.grade]
        row_cells[4].text = applicant.major
        row_cells[5].text = applicant.admitted_department.name if applicant.admitted_department else ''

    # 添加总结信息
    doc.add_paragraph(f'\n总录取人数：{admissions.count()}人', style='Intense Quote')
    doc.add_paragraph(f'导出时间：{timezone.now().strftime("%Y年%m月%d日 %H:%M")}')

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 返回响应
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response[
        'Content-Disposition'] = f'attachment; filename="admissions_{dept_name}_{timezone.now().strftime("%Y%m%d")}.docx"'

    return response


@admin_required
def admin_feedback(request):
    """反馈信息展示"""
    # 获取有反馈的申请
    feedbacks = Applicant.objects.exclude(feedback='').order_by('-updated_at')

    return render(request, 'admin/feedback.html', {
        'feedbacks': feedbacks
    })


@admin_required
def export_feedback(request):
    """导出反馈信息为Word文档"""
    # 获取有反馈的申请
    feedbacks = Applicant.objects.exclude(feedback='').order_by('-updated_at')

    # 创建Word文档
    doc = Document()
    doc.add_heading('拾光文创招新反馈意见汇总', 0)

    # 填充数据
    for i, applicant in enumerate(feedbacks, 1):
        doc.add_heading(f'申请 {i}', level=1)
        doc.add_paragraph(f'姓名：{applicant.name}')
        doc.add_paragraph(f'学号：{applicant.student_id}')
        doc.add_paragraph(f'志愿部门：{applicant.first_choice.name}')
        if applicant.second_choice:
            doc.add_paragraph(f'第二志愿：{applicant.second_choice.name}')
        doc.add_paragraph(f'面试结果：{"录取" if applicant.admission_result else "未录取"}')
        if applicant.admitted_department:
            doc.add_paragraph(f'录取部门：{applicant.admitted_department.name}')

        doc.add_heading('反馈意见：', level=2)
        doc.add_paragraph(applicant.feedback, style='List Bullet')
        doc.add_page_break()

    # 添加总结信息
    doc.add_paragraph(f'总反馈数：{feedbacks.count()}条', style='Intense Quote')
    doc.add_paragraph(f'导出时间：{timezone.now().strftime("%Y年%m月%d日 %H:%M")}')

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 返回响应
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response[
        'Content-Disposition'] = f'attachment; filename="feedback_summary_{timezone.now().strftime("%Y%m%d")}.docx"'

    return response
